"""
Resume Parser Service
Extracts structured profile data from PDF/DOCX uploads.
"""

import io
import re
import uuid
from typing import Any, Dict, List, Optional, Tuple


class ResumeParserService:
    """Parse resume files into structured master profile JSON."""

    # Section headers: strip decoration then match keywords
    SECTION_KEYWORDS = {
        'summary': (
            'summary', 'profile', 'objective', 'about me', 'about', 'professional summary',
            'career summary', 'executive summary',
        ),
        'experience': (
            'experience', 'work experience', 'work history', 'employment',
            'employment history', 'professional experience', 'career history',
            'relevant experience', 'professional background',
        ),
        'education': (
            'education', 'academic background', 'academics', 'qualifications',
        ),
        'skills': (
            'skills', 'technical skills', 'core competencies', 'competencies',
            'technologies', 'technical proficiencies', 'areas of expertise',
            'expertise', 'tools', 'technical expertise',
        ),
        'projects': (
            'projects', 'personal projects', 'selected projects', 'key projects',
        ),
        'certifications': (
            'certifications', 'certificates', 'licenses', 'credentials',
        ),
    }

    BULLET_PATTERN = re.compile(r'^[\u2022\u2023\u25E6\u2043\u25AA\u25CF\-\*•◦▪]\s*')
    ENUM_BULLET_PATTERN = re.compile(r'^\d+[\.\)]\s+')

    DATE_RANGE_PATTERN = re.compile(
        r'('
        r'(?:\d{1,2}/\d{2,4}|\d{1,2}-\d{2,4}|\w+\.?\s+\d{4}|\d{4})'
        r')\s*'
        r'(?:[-–—]|\s+to\s+)\s*'
        r'('
        r'(?:\d{1,2}/\d{2,4}|\d{1,2}-\d{2,4}|\w+\.?\s+\d{4}|\d{4}|present|current|now)'
        r')',
        re.I,
    )
    STANDALONE_DATE_LINE = re.compile(
        r'^('
        r'(?:\d{1,2}/\d{2,4}|\w+\.?\s+\d{4}|\d{4})'
        r')\s*'
        r'(?:[-–—]|\s+to\s+)\s*'
        r'('
        r'(?:\d{1,2}/\d{2,4}|\w+\.?\s+\d{4}|\d{4}|present|current|now)'
        r')\s*$',
        re.I,
    )

    EMAIL_PATTERN = re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+')
    PHONE_PATTERN = re.compile(
        r'(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    )
    LINKEDIN_PATTERN = re.compile(
        r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+/?',
        re.I,
    )
    GITHUB_PATTERN = re.compile(r'(?:https?://)?(?:www\.)?github\.com/[\w-]+/?', re.I)
    URL_PATTERN = re.compile(r'https?://[\w./?=#%-]+', re.I)

    DEGREE_PATTERN = re.compile(
        r'\b('
        r'B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|M\.?B\.?A\.?|Ph\.?D\.?|B\.?E\.?|B\.?Tech'
        r'|Associate|Bachelor|Master|Doctor'
        r')\b',
        re.I,
    )

    SKILL_KNOWLEDGE = [
        'python', 'java', 'javascript', 'typescript', 'ruby', 'go', 'golang', 'rust',
        'c++', 'c#', 'php', 'swift', 'kotlin', 'scala', 'matlab',
        'react', 'angular', 'vue', 'node', 'node.js', 'express', 'next.js', 'nextjs',
        'django', 'flask', 'fastapi', 'spring', 'rails', '.net', 'asp.net',
        'sql', 'postgresql', 'postgres', 'mysql', 'mongodb', 'redis', 'elasticsearch',
        'dynamodb', 'sqlite', 'oracle', 'cassandra', 'nosql',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'k8s', 'terraform', 'ansible',
        'jenkins', 'ci/cd', 'github actions', 'gitlab',
        'git', 'linux', 'unix', 'bash', 'shell',
        'agile', 'scrum', 'jira', 'kanban',
        'rest', 'graphql', 'api', 'microservices', 'grpc',
        'machine learning', 'deep learning', 'nlp', 'data science', 'etl', 'spark',
        'html', 'css', 'sass', 'tailwind', 'bootstrap',
        'tableau', 'power bi', 'excel',
    ]

    SKILL_PATTERN = re.compile(
        r'\b(' + '|'.join(re.escape(s) for s in SKILL_KNOWLEDGE if len(s) > 1) + r')\b',
        re.I,
    )

    @classmethod
    def parse_file(
        cls, file_bytes: bytes, filename: str
    ) -> Tuple[Dict[str, Any], float, List[str]]:
        """Parse a resume file.

        Returns ``(profile, confidence, extraction_warnings)``.
        """
        warnings: List[str] = []
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        if ext == 'docx':
            text = cls._extract_docx(file_bytes)
        elif ext == 'pdf':
            text, warnings = cls._extract_pdf(file_bytes)
        elif ext in ('txt', 'text'):
            text = file_bytes.decode('utf-8', errors='replace')
        else:
            raise ValueError(f'Unsupported file type: {ext}. Use PDF, DOCX, or TXT.')

        profile = cls.parse_text(text)
        confidence = cls._estimate_confidence(profile)
        return profile, confidence, warnings

    @classmethod
    def _normalize_text(cls, text: str) -> str:
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = text.replace('\u00a0', ' ')
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    @classmethod
    def _extract_docx(cls, file_bytes: bytes) -> str:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document(io.BytesIO(file_bytes))
        lines: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or '').lower() if para.style else ''
            is_list = 'list' in style_name or 'bullet' in style_name
            if is_list and not cls.BULLET_PATTERN.match(text) and not cls.ENUM_BULLET_PATTERN.match(text):
                text = f'• {text}'
            lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(' | '.join(cells))

        return cls._normalize_text('\n'.join(lines))

    @classmethod
    def _extract_pdf(cls, file_bytes: bytes) -> Tuple[str, List[str]]:
        import pdfplumber

        parts: List[str] = []
        warnings: List[str] = []
        multi_column_pages = 0

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                words = page.extract_words(
                    x_tolerance=2, y_tolerance=2, keep_blank_chars=False
                ) or []
                is_multi = cls._detect_multi_column(words, page.width or 0)
                if is_multi:
                    multi_column_pages += 1
                    page_text = cls._extract_pdf_columns(words, page.width or 0)
                else:
                    page_text = page.extract_text(
                        x_tolerance=2, y_tolerance=2, layout=False
                    ) or ''
                    if not page_text and words:
                        page_text = cls._words_to_lines(words)

                table_text = cls._extract_pdf_tables(page)
                if table_text:
                    if page_text:
                        page_text = f'{page_text}\n{table_text}'
                    else:
                        page_text = table_text
                        warnings.append(
                            f'Page {page_index}: used table extraction because body text was empty.'
                        )

                if page_text:
                    parts.append(page_text)

        if multi_column_pages:
            warnings.append(
                f'Detected multi-column layout on {multi_column_pages} page(s). '
                'Review parsed sections carefully — column order may be imperfect.'
            )

        return cls._normalize_text('\n'.join(parts)), warnings

    @classmethod
    def _detect_multi_column(cls, words: List[dict], page_width: float) -> bool:
        """Heuristic: two dense x-clusters separated by a wide gap."""
        if len(words) < 40 or page_width < 200:
            return False
        mid = page_width / 2
        left = [w for w in words if w.get('x0', 0) < mid - 20]
        right = [w for w in words if w.get('x0', 0) > mid + 20]
        if len(left) < 15 or len(right) < 15:
            return False
        # Require a sparse middle band (gutter)
        gutter = [
            w for w in words
            if mid - 25 <= w.get('x0', 0) <= mid + 25
        ]
        return len(gutter) < max(8, int(0.08 * len(words)))

    @classmethod
    def _extract_pdf_columns(cls, words: List[dict], page_width: float) -> str:
        """Read left column top-to-bottom, then right column."""
        mid = page_width / 2
        left = [w for w in words if w.get('x0', 0) < mid]
        right = [w for w in words if w.get('x0', 0) >= mid]
        chunks = []
        for column in (left, right):
            if column:
                chunks.append(cls._words_to_lines(column))
        return '\n'.join(chunks)

    @classmethod
    def _extract_pdf_tables(cls, page) -> str:
        """Fallback: flatten table cells into readable lines."""
        try:
            tables = page.extract_tables() or []
        except Exception:
            return ''
        lines: List[str] = []
        for table in tables:
            for row in table or []:
                cells = [
                    re.sub(r'\s+', ' ', (cell or '').strip())
                    for cell in row
                    if cell and str(cell).strip()
                ]
                if cells:
                    lines.append(' | '.join(cells))
        return '\n'.join(lines)

    @classmethod
    def _words_to_lines(cls, words: List[dict], line_tolerance: float = 3.0) -> str:
        if not words:
            return ''
        sorted_words = sorted(words, key=lambda w: (round(w['top'], 1), w['x0']))
        lines: List[str] = []
        current_line: List[str] = []
        current_top: Optional[float] = None

        for word in sorted_words:
            top = round(word['top'], 1)
            if current_top is None or abs(top - current_top) <= line_tolerance:
                current_line.append(word['text'])
                current_top = top if current_top is None else current_top
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word['text']]
                current_top = top
        if current_line:
            lines.append(' '.join(current_line))
        return '\n'.join(lines)

    @classmethod
    def _clean_line(cls, line: str) -> str:
        line = line.strip()
        line = re.sub(r'^#+\s*', '', line)
        line = re.sub(r'^\d+[\.\)]\s+', '', line)
        line = line.rstrip(':').strip()
        return line

    @classmethod
    def _is_section_header(cls, line: str) -> Optional[str]:
        cleaned = cls._clean_line(line).lower()
        if len(cleaned) > 50:
            return None
        for section, keywords in cls.SECTION_KEYWORDS.items():
            for kw in keywords:
                if cleaned == kw or cleaned == kw + 's':
                    return section
                if cleaned.rstrip(':').strip() == kw:
                    return section
                if len(cleaned) <= 40 and re.match(rf'^{re.escape(kw)}\b', cleaned):
                    remainder = cleaned[len(kw):].strip(' :.-')
                    if not remainder or remainder in ('s', 'and', '&'):
                        return section
        if cleaned.isupper() and len(cleaned.split()) <= 5:
            for section, keywords in cls.SECTION_KEYWORDS.items():
                if any(kw in cleaned for kw in keywords):
                    return section
        return None

    @classmethod
    def parse_text(cls, text: str) -> Dict[str, Any]:
        text = cls._normalize_text(text)
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        contact = cls._extract_contact(lines, text)
        sections = cls._split_sections(lines)
        experience = cls._parse_experience(sections.get('experience', []))
        education = cls._parse_education(sections.get('education', []))
        skills = cls._parse_skills(sections.get('skills', []), sections.get('certifications', []))
        summary_variants = cls._parse_summary(sections.get('summary', []))
        projects = cls._parse_projects(sections.get('projects', []))
        headline = cls._guess_headline(lines, contact)

        if not skills.get('technical'):
            skills['technical'] = cls._infer_skills_from_experience(experience)
        if skills.get('certifications'):
            skills['technical'] = list(dict.fromkeys(
                skills['technical'] + skills['certifications']
            ))
            skills['certifications'] = []

        return {
            'contact': contact,
            'headline': headline,
            'summary_variants': summary_variants,
            'experience': experience,
            'education': education,
            'skills': skills,
            'projects': projects,
        }

    @classmethod
    def _extract_contact(cls, lines: List[str], full_text: str) -> Dict[str, str]:
        contact = {
            'name': '',
            'email': '',
            'phone': '',
            'location': '',
            'linkedin': '',
            'website': '',
        }

        email = cls.EMAIL_PATTERN.search(full_text)
        if email:
            contact['email'] = email.group(0)

        phone = cls.PHONE_PATTERN.search(full_text)
        if phone:
            contact['phone'] = phone.group(0).strip()

        linkedin = cls.LINKEDIN_PATTERN.search(full_text)
        if linkedin:
            contact['linkedin'] = linkedin.group(0)

        for line in lines[:12]:
            if cls._is_section_header(line):
                break
            if not contact['name'] and cls._looks_like_name(line):
                contact['name'] = line
                continue
            if not contact['location'] and cls._looks_like_location(line):
                contact['location'] = line

        if not contact['name'] and lines:
            first = lines[0]
            if not cls.EMAIL_PATTERN.search(first) and not cls.PHONE_PATTERN.search(first):
                if len(first) < 60 and not cls._is_section_header(first):
                    contact['name'] = first

        return contact

    @classmethod
    def _looks_like_name(cls, line: str) -> bool:
        if '@' in line or cls.PHONE_PATTERN.search(line) or 'linkedin' in line.lower():
            return False
        if cls.URL_PATTERN.search(line):
            return False
        if cls._is_section_header(line):
            return False
        words = line.split()
        if not 2 <= len(words) <= 4:
            return False
        if any(len(w) > 20 for w in words):
            return False
        return all(w[0].isupper() for w in words if w.isalpha())

    @classmethod
    def _looks_like_location(cls, line: str) -> bool:
        if '@' in line or cls.PHONE_PATTERN.search(line):
            return False
        if len(line) > 80:
            return False
        return ',' in line or re.search(
            r'\b(?:remote|hybrid|city|area)\b', line, re.I
        ) or re.search(r'[A-Z]{2}\s+\d{5}', line)

    @classmethod
    def _split_sections(cls, lines: List[str]) -> Dict[str, List[str]]:
        sections: Dict[str, List[str]] = {'header': []}
        current = 'header'
        for line in lines:
            section = cls._is_section_header(line)
            if section:
                current = section
                sections.setdefault(current, [])
            else:
                sections.setdefault(current, []).append(line)
        return sections

    @classmethod
    def _parse_summary(cls, lines: List[str]) -> List[Dict[str, Any]]:
        if not lines:
            return []
        paragraphs: List[str] = []
        buf: List[str] = []
        for line in lines:
            if cls._is_bullet(line) or cls.DATE_RANGE_PATTERN.search(line):
                if buf:
                    paragraphs.append(' '.join(buf))
                    buf = []
            else:
                buf.append(line)
        if buf:
            paragraphs.append(' '.join(buf))
        text = max(paragraphs, key=len) if paragraphs else ' '.join(lines)
        return [{'id': str(uuid.uuid4()), 'text': text.strip(), 'tags': []}]

    @classmethod
    def _is_bullet(cls, line: str) -> bool:
        return bool(cls.BULLET_PATTERN.match(line) or cls.ENUM_BULLET_PATTERN.match(line))

    @classmethod
    def _strip_bullet(cls, line: str) -> str:
        line = cls.BULLET_PATTERN.sub('', line)
        line = cls.ENUM_BULLET_PATTERN.sub('', line)
        return line.strip()

    @classmethod
    def _is_job_header_line(cls, line: str) -> bool:
        if cls._is_bullet(line):
            return False
        if cls.STANDALONE_DATE_LINE.match(line.strip()):
            return True
        if cls.DATE_RANGE_PATTERN.search(line) and len(line) < 150:
            return True
        if re.search(r'\b(at|@)\b', line, re.I) and len(line) < 120:
            return True
        if '|' in line and len(line) < 120:
            return True
        if cls._looks_like_title_company_comma(line):
            return True
        if cls._looks_like_job_title(line):
            return True
        return False

    @classmethod
    def _looks_like_title_company_comma(cls, line: str) -> bool:
        if ',' not in line or len(line) > 100:
            return False
        parts = [p.strip() for p in line.split(',', 1)]
        return len(parts) == 2 and all(len(p) > 1 for p in parts)

    @classmethod
    def _looks_like_job_title(cls, line: str) -> bool:
        if len(line) > 80 or cls._is_bullet(line):
            return False
        if cls.DATE_RANGE_PATTERN.search(line):
            return False
        job_words = (
            'engineer', 'developer', 'manager', 'analyst', 'architect', 'consultant',
            'designer', 'director', 'lead', 'intern', 'specialist', 'administrator',
            'coordinator', 'associate', 'scientist', 'programmer',
        )
        lower = line.lower()
        return any(w in lower for w in job_words)

    @classmethod
    def _parse_experience(cls, lines: List[str]) -> List[Dict[str, Any]]:
        if not lines:
            return []

        entries: List[Dict[str, Any]] = []
        current: Optional[Dict[str, Any]] = None
        pending_title: Optional[str] = None
        pending_company: Optional[str] = None

        def flush_current() -> None:
            nonlocal current
            if current and (current.get('title') or current.get('company') or current.get('bullets')):
                entries.append(current)
            current = None

        def start_entry(title: str, company: str, start: str = '', end: str = '') -> None:
            nonlocal current, pending_title, pending_company
            flush_current()
            current = cls._new_experience_entry(title, company, start, end)
            pending_title = None
            pending_company = None

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            if cls.STANDALONE_DATE_LINE.match(stripped):
                start, end = cls._extract_dates(stripped)
                if pending_title:
                    company = pending_company or ''
                    start_entry(pending_title, company, start, end)
                elif current:
                    current['start'] = start or current.get('start', '')
                    current['end'] = end or current.get('end', '')
                continue

            if cls._is_bullet(line):
                if not current and pending_title:
                    start_entry(pending_title, pending_company or '')
                if current:
                    bullet_text = cls._strip_bullet(line)
                    if bullet_text:
                        current['bullets'].append(cls._make_bullet(bullet_text))
                continue

            if cls._is_job_header_line(stripped):
                title, company, start, end = cls._parse_role_line(stripped)
                if start or end:
                    start_entry(title, company, start, end)
                elif company:
                    start_entry(title, company)
                elif title and cls._looks_like_job_title(title):
                    flush_current()
                    pending_title = title
                    pending_company = None
                elif title:
                    start_entry(title, company)
                continue

            if pending_title and not pending_company and len(stripped) < 60:
                if not cls._looks_like_job_title(stripped) and not cls._is_bullet(stripped):
                    pending_company = stripped
                    continue

            if current:
                if len(stripped) < 250:
                    current['bullets'].append(cls._make_bullet(stripped))
            elif cls._looks_like_job_title(stripped):
                pending_title = stripped
            else:
                pending_company = stripped

        flush_current()
        for entry in entries:
            entry['bullets'] = [b for b in entry.get('bullets', []) if b.get('text', '').strip()]
        return [e for e in entries if e.get('title') or e.get('company') or e.get('bullets')]

    @classmethod
    def _new_experience_entry(cls, title: str, company: str, start: str, end: str) -> Dict[str, Any]:
        return {
            'id': str(uuid.uuid4()),
            'company': company,
            'title': title,
            'start': start,
            'end': end,
            'location': '',
            'bullets': [],
        }

    @classmethod
    def make_bullet(cls, text: str) -> Dict[str, Any]:
        return cls._make_bullet(text)

    @classmethod
    def _make_bullet(cls, text: str) -> Dict[str, Any]:
        return {
            'id': str(uuid.uuid4()),
            'text': text,
            'skills': cls._extract_skills_from_text(text),
            'metrics': bool(re.search(r'\d+%?|\$[\d,]+|\d+[kKmM]?\+?', text)),
        }

    @classmethod
    def _extract_dates(cls, text: str) -> Tuple[str, str]:
        match = cls.DATE_RANGE_PATTERN.search(text)
        if match:
            return match.group(1).strip(), match.group(2).strip()
        return '', ''

    @classmethod
    def _split_title_company(cls, line: str) -> Tuple[str, str]:
        date_match = cls.DATE_RANGE_PATTERN.search(line)
        if date_match:
            line = line[:date_match.start()].strip(' ,|–—-')

        for sep in (r'\s+at\s+', r'\s+@\s+', r'\s*\|\s*', r'\s+-\s+', r',\s*'):
            parts = re.split(sep, line, maxsplit=1, flags=re.I)
            if len(parts) == 2:
                left, right = parts[0].strip(), parts[1].strip()
                if cls._looks_like_company(right):
                    return left, right
                if cls._looks_like_company(left):
                    return right, left
                return left, right

        return line.strip(), ''

    @classmethod
    def _looks_like_company(cls, text: str) -> bool:
        corp_markers = ('inc', 'llc', 'ltd', 'corp', 'co.', 'company', 'group', 'labs', 'technologies')
        lower = text.lower()
        return any(m in lower for m in corp_markers)

    @classmethod
    def _parse_role_line(cls, line: str) -> Tuple[str, str, str, str]:
        start, end = cls._extract_dates(line)
        title, company = cls._split_title_company(line)
        return title, company, start, end

    @classmethod
    def _parse_education(cls, lines: List[str]) -> List[Dict[str, Any]]:
        entries = []
        for line in lines:
            if cls._is_bullet(line):
                line = cls._strip_bullet(line)
            start, end = cls._extract_dates(line)
            degree_match = cls.DEGREE_PATTERN.search(line)
            degree = degree_match.group(0) if degree_match else ''
            institution = line
            field = ''
            if degree:
                institution = cls.DEGREE_PATTERN.sub('', line).strip(' ,-|–—')
            if ' in ' in institution.lower():
                parts = re.split(r'\s+in\s+', institution, maxsplit=1, flags=re.I)
                if len(parts) == 2:
                    field = parts[1].strip(' ,|–—')
                    institution = parts[0].strip(' ,|–—')
            institution = re.sub(r'\s*[-–|]\s*', ' - ', institution).strip(' -')
            entries.append({
                'id': str(uuid.uuid4()),
                'institution': institution,
                'degree': degree,
                'field': field,
                'start': start,
                'end': end,
            })
        return entries

    @classmethod
    def _parse_skills(
        cls,
        skill_lines: List[str],
        cert_lines: Optional[List[str]] = None,
    ) -> Dict[str, List[str]]:
        technical: List[str] = []
        certifications: List[str] = []

        all_lines = list(skill_lines) + list(cert_lines or [])
        for line in all_lines:
            if cls._is_bullet(line):
                line = cls._strip_bullet(line)
            if ':' in line and len(line.split(':', 1)[0]) < 30:
                _, values = line.split(':', 1)
                technical.extend(cls._split_skill_tokens(values))
            else:
                technical.extend(cls._split_skill_tokens(line))

        technical = [t for t in dict.fromkeys(technical) if len(t) > 1]
        return {'technical': technical, 'certifications': certifications}

    @classmethod
    def _split_skill_tokens(cls, text: str) -> List[str]:
        tokens = re.split(r'[,;|•/\n]|(?:\s{2,})', text)
        result = []
        for token in tokens:
            token = token.strip(' .-•')
            if token and len(token) < 50:
                result.append(token)
        return result

    @classmethod
    def _infer_skills_from_experience(cls, experience: List[Dict[str, Any]]) -> List[str]:
        found: List[str] = []
        for entry in experience:
            for bullet in entry.get('bullets', []):
                found.extend(bullet.get('skills', []))
        return list(dict.fromkeys(found))

    @classmethod
    def _parse_projects(cls, lines: List[str]) -> List[Dict[str, Any]]:
        projects = []
        for line in lines:
            if cls._is_bullet(line):
                line = cls._strip_bullet(line)
            projects.append({
                'id': str(uuid.uuid4()),
                'name': line,
                'description': '',
                'skills': cls._extract_skills_from_text(line),
            })
        return projects

    @classmethod
    def _guess_headline(cls, lines: List[str], contact: Dict[str, str]) -> str:
        name = contact.get('name', '')
        for line in lines[1:8]:
            if line == name:
                continue
            if cls._is_section_header(line):
                break
            if '@' in line or cls.PHONE_PATTERN.search(line) or cls.URL_PATTERN.search(line):
                continue
            if cls._looks_like_location(line):
                continue
            if 5 < len(line) < 90:
                return line
        return ''

    @classmethod
    def _extract_skills_from_text(cls, text: str) -> List[str]:
        return list(dict.fromkeys(m.group(0).lower() for m in cls.SKILL_PATTERN.finditer(text)))

    @classmethod
    def _estimate_confidence(cls, profile: Dict[str, Any]) -> float:
        score = 0.0
        contact = profile.get('contact', {})

        if contact.get('name'):
            score += 12
        if contact.get('email'):
            score += 18
        if contact.get('phone'):
            score += 5
        if contact.get('location'):
            score += 3
        if contact.get('linkedin'):
            score += 2

        experience = profile.get('experience', [])
        if experience:
            score += 20
            score += min(len(experience) * 4, 12)
            bullet_count = sum(len(e.get('bullets', [])) for e in experience)
            score += min(bullet_count * 2, 10)
            dated = sum(1 for e in experience if e.get('start') or e.get('end'))
            score += min(dated * 3, 9)

        education = profile.get('education', [])
        if education:
            score += min(8 + len(education) * 3, 14)

        skills = profile.get('skills', {}).get('technical', [])
        if skills:
            score += min(8 + len(skills), 15)

        if profile.get('summary_variants'):
            score += 5
        if profile.get('headline'):
            score += 3

        return min(round(score, 1), 100.0)

    @classmethod
    def get_parse_diagnostics(
        cls,
        profile: Dict[str, Any],
        extraction_warnings: Optional[List[str]] = None,
    ) -> List[str]:
        """Human-readable hints about what was extracted."""
        hints = list(extraction_warnings or [])
        contact = profile.get('contact', {})
        if not contact.get('email'):
            hints.append('Email not detected — add it to the contact block.')
        if not contact.get('phone'):
            hints.append('Phone number not detected.')
        exp = profile.get('experience', [])
        if not exp:
            hints.append('No work experience found — use headings like "Experience" or "Work History".')
        else:
            hints.append(f'Found {len(exp)} job(s).')
            for i, e in enumerate(exp, 1):
                bullets = len(e.get('bullets', []))
                if bullets == 0:
                    hints.append(f'Job {i} ({e.get("title", "?")}) has no bullets — add • or - prefixes.')
        if not profile.get('skills', {}).get('technical'):
            hints.append('No skills section found — skills may be inferred from experience bullets.')
        if not profile.get('education'):
            hints.append('No education section detected.')
        return hints

    @classmethod
    def validate_profile(cls, profile: Dict[str, Any]) -> List[str]:
        errors = []
        contact = profile.get('contact', {})
        if not contact.get('name'):
            errors.append('Contact name is required.')
        if not contact.get('email'):
            errors.append('Contact email is required.')
        if not profile.get('experience'):
            errors.append('At least one experience entry is required.')
        return errors


resume_parser_service = ResumeParserService()
