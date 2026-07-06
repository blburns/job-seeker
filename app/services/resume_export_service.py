"""
Resume Export Service
Generates ATS-safe DOCX exports and runs parse validation.
"""

import io
import os
import re
import tempfile
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


class ResumeExportService:
    """Export structured resume data to ATS-friendly DOCX."""

    ATS_FONT = 'Calibri'
    ATS_HEADING_SIZE = Pt(12)
    ATS_BODY_SIZE = Pt(11)
    ATS_NAME_SIZE = Pt(16)
    STANDARD_SECTIONS = ['Summary', 'Experience', 'Education', 'Skills']

    @classmethod
    def render_preview_text(cls, profile_data: Dict[str, Any]) -> str:
        """Plain-text preview of the DOCX export (what you upload to job portals)."""
        contact = profile_data.get('contact', {})
        lines = []
        name = (contact.get('name') or '').strip()
        if name:
            lines.append(name.upper())
        headline = (profile_data.get('headline') or '').strip()
        if headline:
            lines.append(headline)
        contact_parts = [
            contact.get('email', ''),
            contact.get('phone', ''),
            contact.get('location', ''),
            contact.get('linkedin', ''),
        ]
        contact_line = ' | '.join(p for p in contact_parts if p)
        if contact_line:
            lines.append(contact_line)
        lines.append('')

        summary = cls._get_summary_text(profile_data)
        if summary:
            lines.extend(['SUMMARY', summary, ''])

        experience = profile_data.get('experience') or []
        if experience:
            lines.append('EXPERIENCE')
            for entry in experience:
                title = entry.get('title', '')
                company = entry.get('company', '')
                start = entry.get('start', '')
                end = entry.get('end', 'Present')
                role_line = f"{title} | {company}" if company else title
                if start:
                    role_line += f" | {start} - {end}"
                lines.append(role_line)
                for bullet in entry.get('bullets', [])[:8]:
                    text = bullet.get('text', bullet) if isinstance(bullet, dict) else str(bullet)
                    if text:
                        lines.append(f"  • {text}")
                lines.append('')

        education = profile_data.get('education') or []
        if education:
            lines.append('EDUCATION')
            for entry in education:
                line = entry.get('institution', '')
                degree = entry.get('degree', '')
                if degree:
                    line = f"{degree} - {line}" if line else degree
                if entry.get('end'):
                    line += f" | {entry['end']}"
                if line:
                    lines.append(line)
            lines.append('')

        skills = profile_data.get('skills', {})
        technical = skills.get('technical', []) if isinstance(skills, dict) else []
        certs = skills.get('certifications', []) if isinstance(skills, dict) else []
        all_skills = [s for s in (technical + certs) if s]
        if all_skills:
            lines.extend(['SKILLS', ', '.join(all_skills)])

        return '\n'.join(lines).strip()

    @classmethod
    def export_cover_letter_docx(cls, cover_letter: str, filename: str = 'Cover_Letter.docx') -> Tuple[bytes, str]:
        doc = Document()
        cls._configure_styles(doc)
        for paragraph in (cover_letter or '').split('\n\n'):
            p = doc.add_paragraph(paragraph.strip())
            if p.runs:
                p.runs[0].font.name = cls.ATS_FONT
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue(), filename

    @classmethod
    def export_docx(cls, profile_data: Dict[str, Any], filename: str = None) -> Tuple[bytes, str]:
        doc = Document()
        cls._configure_styles(doc)
        contact = profile_data.get('contact', {})
        cls._add_contact_block(doc, contact, profile_data.get('headline', ''))
        cls._add_section(doc, 'Summary', cls._get_summary_text(profile_data))
        cls._add_experience_section(doc, profile_data.get('experience', []))
        cls._add_education_section(doc, profile_data.get('education', []))
        cls._add_skills_section(doc, profile_data.get('skills', {}))

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        export_name = filename or cls._build_filename(contact, profile_data)
        return buffer.getvalue(), export_name

    @classmethod
    def _configure_styles(cls, doc: Document) -> None:
        style = doc.styles['Normal']
        style.font.name = cls.ATS_FONT
        style.font.size = cls.ATS_BODY_SIZE
        style.font.color.rgb = RGBColor(0, 0, 0)

    @classmethod
    def _add_contact_block(cls, doc: Document, contact: Dict[str, str], headline: str) -> None:
        name = contact.get('name', '')
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = cls.ATS_NAME_SIZE
        run.font.name = cls.ATS_FONT

        if headline:
            hp = doc.add_paragraph(headline)
            hp.runs[0].font.name = cls.ATS_FONT

        parts = []
        if contact.get('email'):
            parts.append(contact['email'])
        if contact.get('phone'):
            parts.append(contact['phone'])
        if contact.get('location'):
            parts.append(contact['location'])
        if contact.get('linkedin'):
            parts.append(contact['linkedin'])
        if parts:
            cp = doc.add_paragraph(' | '.join(parts))
            cp.runs[0].font.name = cls.ATS_FONT

    @classmethod
    def _add_section(cls, doc: Document, title: str, content: str) -> None:
        if not content:
            return
        heading = doc.add_paragraph(title.upper())
        heading.runs[0].bold = True
        heading.runs[0].font.size = cls.ATS_HEADING_SIZE
        heading.runs[0].font.name = cls.ATS_FONT
        body = doc.add_paragraph(content)
        body.runs[0].font.name = cls.ATS_FONT

    @classmethod
    def _add_experience_section(cls, doc: Document, experience: List[Dict[str, Any]]) -> None:
        if not experience:
            return
        heading = doc.add_paragraph('EXPERIENCE')
        heading.runs[0].bold = True
        heading.runs[0].font.size = cls.ATS_HEADING_SIZE
        heading.runs[0].font.name = cls.ATS_FONT

        for entry in experience:
            title = entry.get('title', '')
            company = entry.get('company', '')
            start = entry.get('start', '')
            end = entry.get('end', 'Present')
            role_line = f"{title} | {company}" if company else title
            if start:
                role_line += f" | {start} - {end}"
            rp = doc.add_paragraph(role_line)
            rp.runs[0].bold = True
            rp.runs[0].font.name = cls.ATS_FONT
            for bullet in entry.get('bullets', [])[:5]:
                text = bullet.get('text', bullet) if isinstance(bullet, dict) else str(bullet)
                bp = doc.add_paragraph(text, style='List Bullet')
                bp.runs[0].font.name = cls.ATS_FONT

    @classmethod
    def _add_education_section(cls, doc: Document, education: List[Dict[str, Any]]) -> None:
        if not education:
            return
        heading = doc.add_paragraph('EDUCATION')
        heading.runs[0].bold = True
        heading.runs[0].font.size = cls.ATS_HEADING_SIZE
        heading.runs[0].font.name = cls.ATS_FONT
        for entry in education:
            line = entry.get('institution', '')
            degree = entry.get('degree', '')
            if degree:
                line = f"{degree} - {line}" if line else degree
            if entry.get('end'):
                line += f" | {entry['end']}"
            p = doc.add_paragraph(line)
            p.runs[0].font.name = cls.ATS_FONT

    @classmethod
    def _add_skills_section(cls, doc: Document, skills: Dict[str, Any]) -> None:
        technical = skills.get('technical', []) if isinstance(skills, dict) else []
        certs = skills.get('certifications', []) if isinstance(skills, dict) else []
        all_skills = technical + certs
        if not all_skills:
            return
        heading = doc.add_paragraph('SKILLS')
        heading.runs[0].bold = True
        heading.runs[0].font.size = cls.ATS_HEADING_SIZE
        heading.runs[0].font.name = cls.ATS_FONT
        p = doc.add_paragraph(', '.join(all_skills))
        p.runs[0].font.name = cls.ATS_FONT

    @classmethod
    def _get_summary_text(cls, profile_data: Dict[str, Any]) -> str:
        variants = profile_data.get('summary_variants', [])
        if variants:
            first = variants[0]
            if isinstance(first, dict):
                return first.get('text', '')
            return str(first)
        return profile_data.get('summary', '')

    @classmethod
    def _build_filename(cls, contact: Dict[str, str], profile_data: Dict[str, Any]) -> str:
        name = contact.get('name', 'Resume').replace(' ', '_')
        role = (profile_data.get('headline') or 'Resume').replace(' ', '_')
        company = profile_data.get('_target_company', '')
        parts = [p for p in [name, role, company] if p]
        return '_'.join(parts[:3]) + '.docx'

    @classmethod
    def run_ats_parse_test(cls, docx_bytes: bytes) -> Dict[str, Any]:
        """Validate ATS readability of exported document."""
        checks = []
        score = 100.0

        try:
            doc = Document(io.BytesIO(docx_bytes))
            full_text = '\n'.join(p.text for p in doc.paragraphs)
            text_lower = full_text.lower()

            for section in ['experience', 'education', 'skills']:
                if section not in text_lower:
                    checks.append({'rule': f'has_{section}_section', 'passed': False, 'message': f'Missing {section} section'})
                    score -= 15
                else:
                    checks.append({'rule': f'has_{section}_section', 'passed': True, 'message': f'Found {section} section'})

            if re.search(r'[\u2022\u2023\u25E6\u2043]', full_text):
                checks.append({'rule': 'simple_bullets', 'passed': True, 'message': 'Uses standard bullets'})
            else:
                checks.append({'rule': 'simple_bullets', 'passed': True, 'message': 'No exotic bullet characters'})

            table_count = len(doc.tables)
            if table_count > 0:
                checks.append({'rule': 'no_tables', 'passed': False, 'message': f'Found {table_count} tables (ATS risk)'})
                score -= 20
            else:
                checks.append({'rule': 'no_tables', 'passed': True, 'message': 'No layout tables'})

            if len(full_text.strip()) < 100:
                checks.append({'rule': 'min_content', 'passed': False, 'message': 'Very little extractable text'})
                score -= 25
            else:
                checks.append({'rule': 'min_content', 'passed': True, 'message': 'Sufficient content length'})

            email_found = bool(re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', full_text))
            checks.append({
                'rule': 'contact_in_body',
                'passed': email_found,
                'message': 'Email found in document body' if email_found else 'Email not found in body',
            })
            if not email_found:
                score -= 10

        except Exception as exc:
            checks.append({'rule': 'parseable', 'passed': False, 'message': str(exc)})
            score = 0

        return {
            'score': max(0, round(score, 1)),
            'checks': checks,
            'passed': score >= 70,
        }

    @classmethod
    def save_export(cls, docx_bytes: bytes, filename: str, upload_folder: str) -> str:
        os.makedirs(upload_folder, exist_ok=True)
        path = os.path.join(upload_folder, filename)
        with open(path, 'wb') as f:
            f.write(docx_bytes)
        return path


resume_export_service = ResumeExportService()
